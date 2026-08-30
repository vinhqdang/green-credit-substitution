"""
Assemble the double-anonymized Elsevier submission package:
  output/Title_Page.docx   (author-identifying info -- kept separate)
  output/Highlights.docx   (3-5 bullets, <=85 chars each)
  output/Manuscript.docx   (fully anonymized main text, tables, figures, references)
"""
import os
import re
import pandas as pd
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.table import WD_TABLE_ALIGNMENT

MANU_DIR = 'manuscript'
DATA_DIR = 'data/processed'
FIG_DIR = 'output/figures'
OUT_DIR = 'output'
os.makedirs(OUT_DIR, exist_ok=True)


def set_base_style(doc):
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)


def add_markdown_paragraphs(doc, md_text):
    """Small markdown->docx converter for our own hand-authored section files:
    #/##/### headings, blank-line-separated paragraphs, '> ' hypothesis callouts,
    and '**bold**' inline spans. Not a general-purpose markdown parser."""
    blocks = re.split(r'\n\s*\n', md_text.strip())
    for block in blocks:
        block = block.strip()
        if not block:
            continue
        if block.startswith('### '):
            doc.add_heading(block[4:].strip(), level=3)
        elif block.startswith('## '):
            doc.add_heading(block[3:].strip(), level=2)
        elif block.startswith('# '):
            doc.add_heading(block[2:].strip(), level=1)
        elif block.startswith('> '):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.4)
            run = p.add_run(re.sub(r'^>\s*', '', block, flags=re.M).replace('\n', ' '))
            run.bold = True
        elif block.lstrip().startswith('- '):
            for line in block.split('\n'):
                doc.add_paragraph(line.lstrip('- ').strip(), style='List Bullet')
        elif re.match(r'^\|.*\|$', block, flags=re.M):
            add_markdown_table(doc, block)
        else:
            text = ' '.join(line.strip() for line in block.split('\n'))
            add_paragraph_with_bold(doc, text)


def add_paragraph_with_bold(doc, text):
    p = doc.add_paragraph()
    parts = re.split(r'(\*\*[^*]+\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            p.add_run(part[2:-2]).bold = True
        else:
            p.add_run(part)


def add_markdown_table(doc, block):
    rows = [r.strip() for r in block.split('\n') if r.strip()]
    rows = [r for r in rows if not re.match(r'^\|[\s\-:|]+\|$', r)]
    cells = [[c.strip() for c in r.strip('|').split('|')] for r in rows]
    table = doc.add_table(rows=len(cells), cols=len(cells[0]))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Light Grid Accent 1'
    for i, row in enumerate(cells):
        for j, val in enumerate(row):
            table.cell(i, j).text = val
            if i == 0:
                for p in table.cell(i, j).paragraphs:
                    for r in p.runs:
                        r.bold = True


def add_csv_table(doc, csv_path, caption, max_rows=None):
    df = pd.read_csv(csv_path)
    if max_rows:
        df = df.head(max_rows)
    doc.add_paragraph(caption).runs[0].bold = True
    table = doc.add_table(rows=len(df) + 1, cols=len(df.columns))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Light Grid Accent 1'
    for j, col in enumerate(df.columns):
        table.cell(0, j).text = str(col)
        for p in table.cell(0, j).paragraphs:
            for r in p.runs:
                r.bold = True
    for i, row in enumerate(df.itertuples(index=False), start=1):
        for j, val in enumerate(row):
            table.cell(i, j).text = '' if pd.isna(val) else (f'{val:.4f}' if isinstance(val, float) else str(val))
    doc.add_paragraph()


def add_figure(doc, img_path, caption, width_inches=6.0):
    if os.path.exists(img_path):
        doc.add_picture(img_path, width=Inches(width_inches))
        cap = doc.add_paragraph(caption)
        cap.runs[0].italic = True
    else:
        doc.add_paragraph(f'[MISSING FIGURE: {img_path}]')


def build_manuscript():
    doc = Document()
    set_base_style(doc)

    with open(os.path.join(MANU_DIR, '00_abstract_frontmatter.md'), encoding='utf-8') as f:
        abstract_only = f.read().split('## Highlights')[0]
    add_markdown_paragraphs(doc, abstract_only)
    doc.add_page_break()

    for fname in ['01_introduction.md', '02_institutional_background.md',
                  '03_literature_hypotheses.md']:
        with open(os.path.join(MANU_DIR, fname), encoding='utf-8') as f:
            add_markdown_paragraphs(doc, f.read())
        doc.add_page_break()

    with open(os.path.join(MANU_DIR, '04_data_methodology.md'), encoding='utf-8') as f:
        add_markdown_paragraphs(doc, f.read())
    doc.add_paragraph()
    add_csv_table(doc, 'data/variable_definitions.csv',
                  'Table 2. Variable definitions and sources.')
    doc.add_page_break()

    with open(os.path.join(MANU_DIR, '05_results.md'), encoding='utf-8') as f:
        add_markdown_paragraphs(doc, f.read())

    doc.add_page_break()
    add_figure(doc, os.path.join(FIG_DIR, 'fig1_world_map.png'),
               'Figure 1. Green practice adoption and SBFN policy membership, by economy.')
    doc.add_page_break()
    add_figure(doc, os.path.join(FIG_DIR, 'fig2_scatter_regquality.png'),
               'Figure 2. Green adoption vs. regulatory quality, by SBFN status.')
    doc.add_page_break()

    add_csv_table(doc, os.path.join(DATA_DIR, 'table1_country_composition.csv'),
                  'Table 1. Sample composition, by economy.')
    doc.add_page_break()
    add_csv_table(doc, os.path.join(DATA_DIR, 'table2_summary_stats.csv'),
                  'Table 3. Summary statistics.')
    doc.add_page_break()
    add_csv_table(doc, os.path.join(DATA_DIR, 'table4_baseline_regressions.csv'),
                  'Table 4. Baseline classical logit regressions (M1-M3).')
    doc.add_page_break()
    add_csv_table(doc, os.path.join(DATA_DIR, 'table5_multilevel.csv'),
                  'Table 5. Bayesian hierarchical model, posterior summary.')
    doc.add_page_break()
    add_csv_table(doc, os.path.join(DATA_DIR, 'table6_causalforest.csv'),
                  'Table 6. Causal forest DML: ATE, CATEs, and feature importances.')
    doc.add_paragraph()
    add_csv_table(doc, os.path.join(DATA_DIR, 'table6b_cate_by_size.csv'),
                  'Table 6b. Causal forest CATEs by firm-size category (follow-up to Table 6\'s '
                  'feature-importance result).')
    doc.add_page_break()
    add_csv_table(doc, os.path.join(DATA_DIR, 'table7_extension_descriptives.csv'),
                  'Table 7. Extension-sample (6-economy) descriptive rates.')
    doc.add_page_break()
    add_csv_table(doc, os.path.join(DATA_DIR, 'table8_robustness.csv'),
                  'Table 8. Additional robustness checks.')
    doc.add_page_break()

    with open(os.path.join(MANU_DIR, '06_discussion_conclusion.md'), encoding='utf-8') as f:
        add_markdown_paragraphs(doc, f.read())
    doc.add_page_break()

    with open(os.path.join(MANU_DIR, '08_appendix.md'), encoding='utf-8') as f:
        add_markdown_paragraphs(doc, f.read())
    doc.add_paragraph()
    add_csv_table(doc, os.path.join(DATA_DIR, 'appendix_a_sbfn_status.csv'),
                  'Table A1. SBFN policy status and regulatory quality, full 47-economy sample.')
    doc.add_page_break()

    with open(os.path.join(MANU_DIR, '09_additional_material.md'), encoding='utf-8') as f:
        add_markdown_paragraphs(doc, f.read())
    doc.add_page_break()

    with open(os.path.join(MANU_DIR, '07_references.md'), encoding='utf-8') as f:
        add_markdown_paragraphs(doc, f.read())

    out_path = os.path.join(OUT_DIR, 'Manuscript.docx')
    doc.save(out_path)
    print(f'Saved {out_path}')


def build_title_page():
    doc = Document()
    set_base_style(doc)
    with open(os.path.join(MANU_DIR, '00_title_page_note.md'), encoding='utf-8') as f:
        add_markdown_paragraphs(doc, f.read())
    out_path = os.path.join(OUT_DIR, 'Title_Page.docx')
    doc.save(out_path)
    print(f'Saved {out_path}')


def build_highlights():
    doc = Document()
    set_base_style(doc)
    with open(os.path.join(MANU_DIR, '00_abstract_frontmatter.md'), encoding='utf-8') as f:
        text = f.read()
    highlights_block = text.split('## Highlights')[1].strip()
    doc.add_heading('Highlights', level=1)
    add_markdown_paragraphs(doc, highlights_block)
    out_path = os.path.join(OUT_DIR, 'Highlights.docx')
    doc.save(out_path)
    print(f'Saved {out_path}')


if __name__ == '__main__':
    build_title_page()
    build_highlights()
    build_manuscript()
